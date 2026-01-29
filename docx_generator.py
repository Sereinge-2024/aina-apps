"""
Générateur DOCX - VERSION FINALE
Tout en gras comme demandé
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import re
from datetime import datetime

class DocxGenerator:
    
    @staticmethod
    def generer_document(document_data, qr_image_bytes):
        """
        Génère le document DOCX - Version finale
        """
        print("=== GÉNÉRATION DOCX (Version Finale) ===")
        
        doc = Document()
        
        # Configuration page
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # ===================================================================
        # 1. QR CODE EN HAUT À DROITE
        # ===================================================================
        print("Ajout QR Code...")
        
        qr_table = doc.add_table(rows=1, cols=2)
        qr_table.autofit = False
        
        left_cell = qr_table.rows[0].cells[0]
        left_para = left_cell.paragraphs[0]
        left_run = left_para.add_run('LABORATOIRE D\'ANALYSE MICROBIOLOGIQUE')
        left_run.bold = True
        left_run.font.size = Pt(10)
        
        right_cell = qr_table.rows[0].cells[1]
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        qr_added = False
        try:
            qr_filename = 'temp_qr.png'
            with open(qr_filename, 'wb') as f:
                f.write(qr_image_bytes)
            
            if os.path.exists(qr_filename):
                right_run = right_para.add_run()
                right_run.add_picture(qr_filename, width=Inches(1.0))
                qr_added = True
                print("✓ QR Code ajouté")
                os.remove(qr_filename)
        except Exception as e:
            print(f"Erreur QR: {e}")
        
        if not qr_added:
            right_para.add_run('[QR Code]')
        
        doc.add_paragraph()
        
        # ===================================================================
        # 2. DATE DU JOUR (Format: Ouagadougou, le 28 janvier 2026)
        # ===================================================================
        print("Ajout date du jour...")
        
        aujourdhui = datetime.now()
        
        mois_fr = {
            1: 'janvier', 2: 'février', 3: 'mars', 4: 'avril',
            5: 'mai', 6: 'juin', 7: 'juillet', 8: 'août',
            9: 'septembre', 10: 'octobre', 11: 'novembre', 12: 'décembre'
        }
        
        date_formatted = f"Ouagadougou, le {aujourdhui.day} {mois_fr[aujourdhui.month]} {aujourdhui.year}"
        
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(date_formatted)
        date_run.italic = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        # ===================================================================
        # 3. TITRE
        # ===================================================================
        print("Ajout titre...")
        titre = doc.add_paragraph()
        titre_run = titre.add_run('RESULTATS D\'ANALYSE MICROBIOLOGIQUE D\'EAU')
        titre_run.bold = True
        titre_run.font.size = Pt(14)
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # ===================================================================
        # 4. INFORMATIONS (TOUT EN GRAS - labels ET valeurs)
        # ===================================================================
        print("Ajout informations...")
        
        # Analyse n° (TOUT EN GRAS)
        para1 = doc.add_paragraph()
        run1 = para1.add_run(f"Analyse n° : {document_data['document_number']}")
        run1.bold = True
        
        # Date de prélèvement et Lieu (TOUT EN GRAS)
        date_prel = document_data.get('date_prelevement', '')
        if hasattr(date_prel, 'strftime'):
            date_prel = date_prel.strftime('%d/%m/%Y')  # Sans heures
        
        para2 = doc.add_paragraph()
        run2 = para2.add_run(f"Date de prélèvement : {date_prel}                    Lieu : {document_data.get('lieu', '')}")
        run2.bold = True
        
        # Date de réception (TOUT EN GRAS)
        date_recep = document_data.get('date_reception', '')
        if hasattr(date_recep, 'strftime'):
            date_recep = date_recep.strftime('%d/%m/%Y')  # Sans heures
        
        para3 = doc.add_paragraph()
        run3 = para3.add_run(f"Date de réception : {date_recep}")
        run3.bold = True
        
        # Préleveur (TOUT EN GRAS)
        para4 = doc.add_paragraph()
        run4 = para4.add_run(f"Identité du préleveur : {document_data.get('identite_preleveur', '')}")
        run4.bold = True
        
        # Demandeur (TOUT EN GRAS)
        para5 = doc.add_paragraph()
        run5 = para5.add_run(f"Identité du demandeur : {document_data.get('identite_demandeur', '')}")
        run5.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # ===================================================================
        # 5. TABLEAU DES PARAMÈTRES
        # ===================================================================
        print("Création tableau...")
        
        import json
        resultats_json = document_data.get('resultats_json', '[]')
        
        try:
            if isinstance(resultats_json, str):
                resultats = json.loads(resultats_json)
            else:
                resultats = resultats_json
        except Exception as e:
            print(f"Erreur JSON: {e}")
            resultats = []
        
        print(f"Nombre d'échantillons: {len(resultats)}")
        
        if not resultats:
            resultats = [{'coliformes_totaux': '0', 'coliformes_fecaux': '0', 'streptocoques_fecaux': '0'}]
        
        for idx, resultat in enumerate(resultats):
            print(f"Tableau échantillon {idx+1}...")
            
            if idx > 0:
                doc.add_paragraph()
            
            # CRÉER LE TABLEAU
            table = doc.add_table(rows=4, cols=5)
            table.style = 'Light Grid Accent 1'
            
            # EN-TÊTE
            headers = ['PARAMETRES', 'Température\net temps', 'Technique\net milieu', 'RESULTATS\nUFC/100ml', 'NORMES\nOMS']
            
            for i, header_text in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header_text
                
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                
                try:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'D9D9D9')
                    cell._element.get_or_add_tcPr().append(shading)
                except:
                    pass
            
            # DONNÉES
            table.rows[1].cells[0].text = 'Coliformes totaux'
            table.rows[1].cells[1].text = '37°C 24h'
            table.rows[1].cells[2].text = 'Filtration membrane'
            table.rows[1].cells[3].text = str(resultat.get('coliformes_totaux', ''))
            table.rows[1].cells[4].text = '0/100 ml'
            
            table.rows[2].cells[0].text = 'Coliformes fécaux'
            table.rows[2].cells[1].text = '44°C 24h'
            table.rows[2].cells[2].text = 'Filtration membrane'
            table.rows[2].cells[3].text = str(resultat.get('coliformes_fecaux', ''))
            table.rows[2].cells[4].text = '0/100 ml'
            
            table.rows[3].cells[0].text = 'Streptocoques fécaux'
            table.rows[3].cells[1].text = '24h'
            table.rows[3].cells[2].text = 'Filtration membrane'
            table.rows[3].cells[3].text = str(resultat.get('streptocoques_fecaux', ''))
            table.rows[3].cells[4].text = '0/100 ml'
            
            print(f"  Résultats: CT={resultat.get('coliformes_totaux')}, CF={resultat.get('coliformes_fecaux')}, SF={resultat.get('streptocoques_fecaux')}")
            
            # Centrer toutes les cellules
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(9)
            
            print(f"  ✓ Tableau créé")
        
        # ===================================================================
        # 6. CONCLUSION
        # ===================================================================
        print("Ajout conclusion...")
        doc.add_paragraph()
        
        conclusion_para = doc.add_paragraph()
        conclusion_run = conclusion_para.add_run('Conclusion : ')
        conclusion_run.bold = True
        
        conclusion_text = document_data.get('conclusion', '')
        conclusion_text = DocxGenerator._clean_html(conclusion_text)
        conclusion_para.add_run(conclusion_text)
        
        # ===================================================================
        # 7. NB
        # ===================================================================
        note_text = document_data.get('note', '')
        if note_text:
            print("Ajout NB...")
            note_para = doc.add_paragraph()
            note_run = note_para.add_run('NB : ')
            note_run.bold = True
            note_text = DocxGenerator._clean_html(note_text)
            note_para.add_run(note_text)
        
        # ===================================================================
        # 8. SIGNATURE À DROITE
        # ===================================================================
        print("Ajout signature...")
        doc.add_paragraph()
        
        titre_sig = doc.add_paragraph()
        titre_sig.add_run(document_data.get('titre_signataire', 'Le Chef du Laboratoire'))
        titre_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        nom_sig = doc.add_paragraph()
        nom_sig.add_run(document_data.get('nom_signataire', ''))
        nom_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # ===================================================================
        # SAUVEGARDER
        # ===================================================================
        print("Sauvegarde...")
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        size = len(docx_file.getvalue())
        print(f"✓ Document généré: {size} bytes")
        print("=== FIN GÉNÉRATION ===")
        
        return docx_file
    
    @staticmethod
    def _clean_html(html_text):
        """Nettoie le HTML"""
        if not html_text:
            return ''
        
        text = re.sub(r'<br\s*/?>', '\n', html_text)
        text = re.sub(r'<p>', '', text)
        text = re.sub(r'</p>', '\n', text)
        text = re.sub(r'<[^>]+>', '', text)
        text = text.replace('&nbsp;', ' ')
        text = text.strip()
        
        return text
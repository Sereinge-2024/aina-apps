"""
Générateur de documents DOCX - FORMAT EXACT
Respecte à 100% la structure du document original
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import tempfile
import re

class DocxGenerator:
    
    @staticmethod
    def generer_document(document_data, qr_image_bytes):
        """
        Génère le document DOCX selon le format EXACT
        """
        doc = Document()
        
        # Configuration page A4
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # ===================================================================
        # 1. EN-TÊTE DE PAGE (avec logo et QR Code)
        # ===================================================================
        DocxGenerator._creer_entete(doc, qr_image_bytes)
        
        # ===================================================================
        # 2. TITRE PRINCIPAL (centré)
        # ===================================================================
        doc.add_paragraph()  # Espace
        
        titre = doc.add_paragraph()
        titre_run = titre.add_run('RESULTATS D\'ANALYSE MICROBIOLOGIQUE D\'EAU')
        titre_run.bold = True
        titre_run.font.size = Pt(14)
        titre_run.font.name = 'Arial'
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espace
        
        # ===================================================================
        # 3. INFORMATIONS (paragraphes simples avec tabulations)
        # ===================================================================
        
        # Numéro d'analyse
        para1 = doc.add_paragraph()
        para1.add_run(f"Analyse n° : {document_data['document_number']}\t\t\t\t")
        
        # Date de prélèvement et Lieu (sur la même ligne)
        date_prel = document_data.get('date_prelevement', '')
        if hasattr(date_prel, 'strftime'):
            date_prel = date_prel.strftime('%d/%m/%Y')
        
        para2 = doc.add_paragraph()
        para2.add_run(f"Date de prélèvement : {date_prel}\t\t\t\tLieu : {document_data.get('lieu', '')}")
        
        # Date de réception
        date_recep = document_data.get('date_reception', '')
        if hasattr(date_recep, 'strftime'):
            date_recep = date_recep.strftime('%d/%m/%Y')
        
        para3 = doc.add_paragraph()
        para3.add_run(f"Date de réception : {date_recep} \t\t\t\t")
        
        # Préleveur
        para4 = doc.add_paragraph()
        para4.add_run(f"Identité du préleveur : {document_data.get('identite_preleveur', '')}\t")
        
        # Demandeur
        para5 = doc.add_paragraph()
        para5.add_run(f"Identité du demandeur : {document_data.get('identite_demandeur', '')}")
        
        doc.add_paragraph()  # Espace avant tableau
        doc.add_paragraph()  # Double espace
        
        # ===================================================================
        # 4. TABLEAUX DES PARAMÈTRES (un par échantillon)
        # ===================================================================
        
        import json
        resultats_json = document_data.get('resultats_json', '[]')
        
        try:
            if isinstance(resultats_json, str):
                resultats = json.loads(resultats_json)
            else:
                resultats = resultats_json
        except:
            resultats = []
        
        for idx, resultat in enumerate(resultats):
            if idx > 0:
                doc.add_paragraph()  # Espace entre tableaux
                doc.add_paragraph()
            
            DocxGenerator._creer_tableau_parametres(doc, resultat)
        
        # ===================================================================
        # 5. CONCLUSION (à gauche, après le tableau)
        # ===================================================================
        doc.add_paragraph()  # Espace après tableau
        
        conclusion_para = doc.add_paragraph()
        conclusion_run = conclusion_para.add_run('Conclusion : ')
        conclusion_run.bold = True
        
        conclusion_text = document_data.get('conclusion', '')
        conclusion_text = DocxGenerator._clean_html(conclusion_text)
        conclusion_para.add_run(conclusion_text)
        conclusion_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # ===================================================================
        # 6. NB (à gauche)
        # ===================================================================
        note_text = document_data.get('note', '')
        if note_text:
            note_para = doc.add_paragraph()
            note_run = note_para.add_run('NB : ')
            note_run.bold = True
            note_text = DocxGenerator._clean_html(note_text)
            note_para.add_run(note_text)
            note_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # ===================================================================
        # 7. SIGNATURE (EN BAS À DROITE)
        # ===================================================================
        doc.add_paragraph()  # Espace
        
        # Titre du signataire
        titre_sig_para = doc.add_paragraph()
        titre_sig_para.add_run(document_data.get('titre_signataire', 'Le Chef du Laboratoire') + ' ')
        titre_sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Espaces pour signature manuscrite
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Nom du signataire
        nom_sig_para = doc.add_paragraph()
        nom_sig_run = nom_sig_para.add_run(document_data.get('nom_signataire', ''))
        nom_sig_run.bold = False
        nom_sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Sauvegarder
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        return docx_file
    
    @staticmethod
    def _creer_entete(doc, qr_image_bytes):
        """
        Crée l'en-tête exactement comme dans le document original
        Structure: Tableau 4 lignes x 2 colonnes
        """
        header = doc.sections[0].header
        
        # Créer le tableau d'en-tête (4 lignes x 2 colonnes)
        header_table = header.add_table(rows=4, cols=2)
        header_table.autofit = False
        
        # Définir les largeurs
        header_table.columns[0].width = Inches(5.5)
        header_table.columns[1].width = Inches(1.5)
        
        # LIGNE 1: Logo (gauche) et QR Code (droite)
        row1_left = header_table.rows[0].cells[0]
        row1_right = header_table.rows[0].cells[1]
        
        # Logo à gauche
        logo_paths = [
            os.path.join('static', 'images', 'image1.jpeg'),
            'static/images/image1.jpeg',
            'image1.jpeg'
        ]
        
        logo_added = False
        for logo_path in logo_paths:
            if os.path.exists(logo_path):
                try:
                    para_logo = row1_left.paragraphs[0]
                    run_logo = para_logo.add_run()
                    run_logo.add_picture(logo_path, width=Inches(2.0))
                    logo_added = True
                    break
                except Exception as e:
                    print(f"Erreur logo: {e}")
        
        if not logo_added:
            # Texte de remplacement
            para_logo = row1_left.paragraphs[0]
            run_logo = para_logo.add_run('LABORATOIRE D\'ANALYSE\nMICROBIOLOGIQUE DE L\'EAU')
            run_logo.font.size = Pt(12)
            run_logo.bold = True
        
        # QR Code à droite
        para_qr = row1_right.paragraphs[0]
        para_qr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_qr:
                tmp_qr.write(qr_image_bytes)
                tmp_qr_path = tmp_qr.name
            
            run_qr = para_qr.add_run()
            run_qr.add_picture(tmp_qr_path, width=Inches(1.2))
            
            os.unlink(tmp_qr_path)
        except Exception as e:
            print(f"Erreur QR: {e}")
        
        # LIGNE 2, 3, 4: Informations de l'entreprise (fusionnées sur toute la largeur)
        # On peut ajouter du texte ici si nécessaire
        # Pour l'instant, on laisse vide comme dans l'original
        
        # Fusionner les cellules des lignes 2-4 pour avoir une zone unique
        # (optionnel selon le besoin)
    
    @staticmethod
    def _creer_tableau_parametres(doc, resultat):
        """
        Crée le tableau des paramètres EXACTEMENT comme l'original
        4 lignes x 5 colonnes
        """
        
        table = doc.add_table(rows=4, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Définir les largeurs approximatives
        widths = [Inches(2.2), Inches(1.3), Inches(1.5), Inches(1.0), Inches(1.5)]
        for i, width in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = width
        
        # ============================================================
        # LIGNE 0: EN-TÊTE (avec fond gris)
        # ============================================================
        headers = [
            '\nPARAMETRES\n',
            '\nTempérature\net temps d\'incubation',
            '\nTechnique et milieu de culture',
            '\nRESULTATS\nUFC/100 ml',
            '\nNORMES DE QUALITE OMS POUR EAU POTABLE'
        ]
        
        for i, header_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            para = cell.paragraphs[0]
            run = para.add_run(header_text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Fond gris clair
            DocxGenerator._set_cell_shading(cell, "D9D9D9")
        
        # ============================================================
        # LIGNE 1: Coliformes totaux
        # ============================================================
        row1 = table.rows[1].cells
        
        DocxGenerator._set_cell_text(row1[0], 
            '° Recherche et dénombrement\n\ndes Coliformes totaux',
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row1[1], 
            '37°C 24h',
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row1[2], 
            'Filtration sur membrane\n\nChromocult agar Coliformes',
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row1[3], 
            str(resultat.get('coliformes_totaux', '')),
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row1[4], 
            '0/100 ml',
            centered=True, size=9)
        
        # ============================================================
        # LIGNE 2: Coliformes fécaux
        # ============================================================
        row2 = table.rows[2].cells
        
        DocxGenerator._set_cell_text(row2[0], 
            '° Recherche et dénombrement\n\ndes Coliformes fécaux',
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row2[1], 
            '44°C 24h',
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row2[2], 
            'Filtration sur membrane\n\nChromocult agar Coliformes',
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row2[3], 
            str(resultat.get('coliformes_fecaux', '')),
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row2[4], 
            '0/100 ml',
            centered=True, size=9)
        
        # ============================================================
        # LIGNE 3: Streptocoques fécaux
        # ============================================================
        row3 = table.rows[3].cells
        
        DocxGenerator._set_cell_text(row3[0], 
            '° Recherche et dénombrement\n\ndes Streptocoques fécaux',
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row3[1], 
            '24h.',
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row3[2], 
            'Filtration sur membrane\n\nChromocult Entérocoques-agar',
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row3[3], 
            str(resultat.get('streptocoques_fecaux', '')),
            centered=True, size=9)
        
        DocxGenerator._set_cell_text(row3[4], 
            '0/100 ml',
            centered=True, size=9)
    
    @staticmethod
    def _set_cell_text(cell, text, centered=False, size=10, bold=False):
        """Définit le texte d'une cellule avec formatage"""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Arial'
        if bold:
            run.bold = True
        if centered:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    @staticmethod
    def _set_cell_shading(cell, color):
        """Définit la couleur de fond d'une cellule"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        tcPr.append(shading)
    
    @staticmethod
    def _clean_html(html_text):
        """Nettoie le HTML pour avoir du texte pur"""
        if not html_text:
            return ''
        
        text = re.sub(r'<br\s*/?>', '\n', html_text)
        text = re.sub(r'<p>', '', text)
        text = re.sub(r'</p>', '\n', text)
        text = re.sub(r'<[^>]+>', '', text)
        text = text.replace('&nbsp;', ' ')
        text = text.strip()
        
        return text
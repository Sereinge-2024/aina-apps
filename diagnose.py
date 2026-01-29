"""
Script de diagnostic pour identifier le problème DOCX
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def test_generation_simple():
    """Test de génération minimale"""
    print("=== TEST 1: Document minimal ===")
    try:
        doc = Document()
        
        # Juste un titre
        para = doc.add_paragraph()
        run = para.add_run('TEST DOCUMENT')
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sauvegarder
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Vérifier la taille
        size = len(output.getvalue())
        print(f"✅ Document minimal créé ({size} bytes)")
        
        # Sauvegarder sur disque pour test
        with open('test_minimal.docx', 'wb') as f:
            f.write(output.getvalue())
        print("✅ Fichier test_minimal.docx créé - Essayez de l'ouvrir dans Word")
        
        return True
        
    except Exception as e:
        print(f"❌ ERREUR: {e}")
        return False

def test_generation_avec_tableau():
    """Test avec tableau"""
    print("\n=== TEST 2: Document avec tableau ===")
    try:
        doc = Document()
        
        para = doc.add_paragraph('Document avec tableau')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tableau simple
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Remplir
        table.rows[0].cells[0].text = 'Col 1'
        table.rows[0].cells[1].text = 'Col 2'
        table.rows[0].cells[2].text = 'Col 3'
        
        table.rows[1].cells[0].text = 'Valeur 1'
        table.rows[1].cells[1].text = 'Valeur 2'
        table.rows[1].cells[2].text = 'Valeur 3'
        
        # Sauvegarder
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        size = len(output.getvalue())
        print(f"✅ Document avec tableau créé ({size} bytes)")
        
        with open('test_tableau.docx', 'wb') as f:
            f.write(output.getvalue())
        print("✅ Fichier test_tableau.docx créé - Essayez de l'ouvrir dans Word")
        
        return True
        
    except Exception as e:
        print(f"❌ ERREUR: {e}")
        return False

def test_generation_complete():
    """Test complet comme l'application"""
    print("\n=== TEST 3: Document complet ===")
    try:
        doc = Document()
        
        # Configuration page
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        
        # Titre
        titre = doc.add_paragraph()
        titre_run = titre.add_run('RESULTATS D\'ANALYSE MICROBIOLOGIQUE D\'EAU')
        titre_run.bold = True
        titre_run.font.size = Pt(14)
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Infos
        doc.add_paragraph("Analyse n° : 001/2026")
        doc.add_paragraph("Date de prélèvement : 28/01/2026                Lieu : Test")
        doc.add_paragraph("Date de réception : 28/01/2026")
        doc.add_paragraph("Identité du préleveur : Test Préleveur")
        doc.add_paragraph("Identité du demandeur : Test Demandeur")
        
        doc.add_paragraph()
        
        # Tableau
        table = doc.add_table(rows=4, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # En-tête
        headers = ['PARAMETRES', 'Température', 'Technique', 'RESULTATS', 'NORMES']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Données
        table.rows[1].cells[0].text = 'Coliformes totaux'
        table.rows[1].cells[1].text = '37°C'
        table.rows[1].cells[2].text = 'Filtration'
        table.rows[1].cells[3].text = '0'
        table.rows[1].cells[4].text = '0/100ml'
        
        # Conclusion
        doc.add_paragraph()
        conclusion = doc.add_paragraph()
        conclusion_run = conclusion.add_run('Conclusion : ')
        conclusion_run.bold = True
        conclusion.add_run('Eau conforme.')
        
        # Signature
        doc.add_paragraph()
        sig = doc.add_paragraph('Le Chef du Laboratoire')
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        nom = doc.add_paragraph('Dr Test')
        nom.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Sauvegarder
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        size = len(output.getvalue())
        print(f"✅ Document complet créé ({size} bytes)")
        
        with open('test_complet.docx', 'wb') as f:
            f.write(output.getvalue())
        print("✅ Fichier test_complet.docx créé - Essayez de l'ouvrir dans Word")
        
        return True
        
    except Exception as e:
        print(f"❌ ERREUR: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    print("DIAGNOSTIC GÉNÉRATION DOCX\n")
    
    test1 = test_generation_simple()
    test2 = test_generation_avec_tableau()
    test3 = test_generation_complete()
    
    print("\n" + "="*50)
    print("RÉSUMÉ:")
    print(f"Test 1 (minimal): {'✅ OK' if test1 else '❌ ÉCHEC'}")
    print(f"Test 2 (tableau): {'✅ OK' if test2 else '❌ ÉCHEC'}")
    print(f"Test 3 (complet): {'✅ OK' if test3 else '❌ ÉCHEC'}")
    print("="*50)
    
    if test3:
        print("\n✅ La génération DOCX fonctionne !")
        print("Ouvrez test_complet.docx dans Word pour vérifier.")
        print("\nSi ce fichier s'ouvre correctement mais pas celui de l'application,")
        print("le problème vient probablement de la gestion des images.")
    else:
        print("\n❌ Problème détecté dans la génération DOCX")
        print("Vérifiez l'installation de python-docx:")
        print("  pip install --upgrade python-docx")